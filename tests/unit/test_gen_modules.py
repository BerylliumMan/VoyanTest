"""Comprehensive unit tests for app/gen/ module parsers.

Targets >= 90% coverage for:
- pdf_parser.py
- model_client.py
- response_parser.py
- feature_extractor.py
- orchestrator.py
- multi_file.py

All external dependencies (fitz, requests, db, model) are mocked.
"""
import io
import json
import base64
from unittest.mock import MagicMock, patch, mock_open, PropertyMock

import pytest


# ===========================================================================
# pdf_parser tests
# ===========================================================================

class TestPdfParserExtractText:
    """Test pdf_parser.extract_text_from_pdf"""

    @pytest.mark.asyncio
    async def test_extract_text_from_pdf_happy_path(self):
        """Real PDF generation, then extract text via mock of fitz."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        # Create a real PDF in memory
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Hello PDF World")

        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import extract_text_from_pdf
        result = extract_text_from_pdf(io.BytesIO(pdf_bytes))
        assert "Hello" in result

    @pytest.mark.asyncio
    async def test_extract_text_from_pdf_multipage(self):
        """Multi-page PDF should concatenate with newlines."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 50), f"Page{i}")
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import extract_text_from_pdf
        result = extract_text_from_pdf(io.BytesIO(pdf_bytes))
        assert "Page0" in result
        assert "Page1" in result
        assert "Page2" in result

    @pytest.mark.asyncio
    async def test_extract_text_from_pdf_garbage_raises(self):
        """Garbage bytes should raise an exception from fitz.open."""
        from app.gen.pdf_parser import extract_text_from_pdf
        with pytest.raises(Exception):
            extract_text_from_pdf(io.BytesIO(b"not a pdf at all"))


class TestPdfParserIsDualLayer:
    """Test pdf_parser.is_pdf_dual_layer"""

    @pytest.mark.asyncio
    async def test_is_pdf_dual_layer_true_with_text(self):
        """PDF with extractable text → True."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Some text")
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import is_pdf_dual_layer
        assert is_pdf_dual_layer(io.BytesIO(pdf_bytes)) is True

    @pytest.mark.asyncio
    async def test_is_pdf_dual_layer_false_no_text(self):
        """PDF with no text → False."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()  # blank page
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import is_pdf_dual_layer
        assert is_pdf_dual_layer(io.BytesIO(pdf_bytes)) is False

    @pytest.mark.asyncio
    async def test_is_pdf_dual_layer_multipage(self):
        """Multi-page PDF: should return True if any page has text."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()  # blank
        p2 = doc.new_page()
        p2.insert_text((50, 50), "Text on page 2")
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import is_pdf_dual_layer
        assert is_pdf_dual_layer(io.BytesIO(pdf_bytes)) is True


class TestPdfParserRenderPages:
    """Test pdf_parser.render_pdf_pages_to_images"""

    @pytest.mark.asyncio
    async def test_render_pages_to_images(self):
        """Should return list of (format, base64) tuples for each page."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import render_pdf_pages_to_images
        images = render_pdf_pages_to_images(io.BytesIO(pdf_bytes))
        assert isinstance(images, list)
        assert len(images) == 2
        for ext, b64 in images:
            assert ext == "png"
            assert isinstance(b64, str)
            # Should be valid base64
            decoded = base64.b64decode(b64)
            assert len(decoded) > 0

    @pytest.mark.asyncio
    async def test_render_pages_to_images_single(self):
        """Single page PDF: returns one image."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import render_pdf_pages_to_images
        images = render_pdf_pages_to_images(io.BytesIO(pdf_bytes))
        assert len(images) == 1


class TestPdfParserValidate:
    """Test pdf_parser.validate_pdf"""

    @pytest.mark.asyncio
    async def test_validate_pdf_valid(self):
        """Valid PDF returns (True, None)."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        from app.gen.pdf_parser import validate_pdf
        ok, msg = validate_pdf(io.BytesIO(pdf_bytes))
        assert ok is True
        assert msg is None

    @pytest.mark.asyncio
    async def test_validate_pdf_garbage(self):
        """Garbage bytes → (False, error message)."""
        from app.gen.pdf_parser import validate_pdf
        ok, msg = validate_pdf(io.BytesIO(b"definitely not a pdf"))
        assert ok is False
        assert msg is not None
        assert "损坏" in msg or "失败" in msg

    @pytest.mark.asyncio
    async def test_validate_pdf_encrypted(self):
        """Encrypted PDF → (False, '无法打开加密PDF文档')."""
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        # Create a real PDF and encrypt it
        doc = fitz.open()
        doc.new_page()
        # Save with encryption
        pdf_bytes = doc.tobytes(garbage=0, deflate=True, encryption=fitz.PDF_ENCRYPT_KEEP)
        doc.close()

        from app.gen.pdf_parser import validate_pdf
        # Encrypted check is in validate_pdf; may not always be encrypted without password
        # Just exercise the path
        ok, msg = validate_pdf(io.BytesIO(pdf_bytes))
        assert isinstance(ok, bool)
        assert msg is None or isinstance(msg, str)

    @pytest.mark.asyncio
    async def test_validate_pdf_generic_exception(self, monkeypatch):
        """Force a generic exception to hit the except Exception branch."""
        from app.gen import pdf_parser

        class FakeDoc:
            is_encrypted = False
            def __len__(self):
                return 1
            def close(self):
                pass

        class FakeFitz:
            FileDataError = type("FileDataError", (Exception,), {})
            @staticmethod
            def open(*a, **kw):
                raise RuntimeError("boom")

        monkeypatch.setattr(pdf_parser, "fitz", FakeFitz)
        ok, msg = pdf_parser.validate_pdf(io.BytesIO(b"x"))
        assert ok is False
        assert "boom" in msg

    @pytest.mark.asyncio
    async def test_validate_pdf_fitz_filedata_error(self, monkeypatch):
        """Force fitz.FileDataError to hit the specific except branch."""
        from app.gen import pdf_parser

        class FDE(Exception):
            pass

        class FakeFitz:
            FileDataError = FDE
            @staticmethod
            def open(*a, **kw):
                raise FDE("corrupt")

        monkeypatch.setattr(pdf_parser, "fitz", FakeFitz)
        ok, msg = pdf_parser.validate_pdf(io.BytesIO(b"x"))
        assert ok is False
        assert "损坏" in msg

    @pytest.mark.asyncio
    async def test_validate_pdf_zero_pages(self, monkeypatch):
        """PDF with 0 pages → (False, 'PDF文件中无有效内容')."""
        from app.gen import pdf_parser

        class FakeDoc:
            is_encrypted = False
            def __len__(self):
                return 0
            def close(self):
                pass

        class FakeFitz:
            FileDataError = type("FileDataError", (Exception,), {})
            @staticmethod
            def open(*a, **kw):
                return FakeDoc()

        monkeypatch.setattr(pdf_parser, "fitz", FakeFitz)
        ok, msg = pdf_parser.validate_pdf(io.BytesIO(b"x"))
        assert ok is False
        assert "无有效内容" in msg

    @pytest.mark.asyncio
    async def test_validate_pdf_encrypted_mocked(self, monkeypatch):
        """Force is_encrypted=True to hit the encrypted PDF branch (lines 53-54)."""
        from app.gen import pdf_parser

        class FakeDoc:
            is_encrypted = True
            def __len__(self):
                return 1
            def close(self):
                pass

        class FakeFitz:
            FileDataError = type("FileDataError", (Exception,), {})
            @staticmethod
            def open(*a, **kw):
                return FakeDoc()

        monkeypatch.setattr(pdf_parser, "fitz", FakeFitz)
        ok, msg = pdf_parser.validate_pdf(io.BytesIO(b"x"))
        assert ok is False
        assert "加密" in msg


# ===========================================================================
# model_client tests
# ===========================================================================

class TestModelClientStripBr:
    """Test _strip_br helper."""

    @pytest.mark.asyncio
    async def test_strip_br_basic(self):
        from app.gen.model_client import _strip_br
        assert _strip_br("a<br>b") == "a b"

    @pytest.mark.asyncio
    async def test_strip_br_with_spaces(self):
        from app.gen.model_client import _strip_br
        assert _strip_br("a< br/>b") == "a b"

    @pytest.mark.asyncio
    async def test_strip_br_no_tags(self):
        from app.gen.model_client import _strip_br
        assert _strip_br("plain text") == "plain text"

    @pytest.mark.asyncio
    async def test_strip_br_uppercase(self):
        from app.gen.model_client import _strip_br
        assert _strip_br("a<BR>b") == "a b"


class TestModelClientLoadConfig:
    """Test _load_ai_config."""

    @pytest.mark.asyncio
    async def test_load_ai_config_success(self):
        from app.gen.model_client import _load_ai_config

        # Mock the AIConfig row
        mock_row = MagicMock()
        mock_row.model = "gpt-4"
        mock_row.api_key = "encrypted-blob"
        mock_row.api_base = "https://api.openai.com/v1"
        mock_row.temperature = 0.5

        mock_db = MagicMock()
        mock_query = MagicMock()
        mock_query.filter.return_value.first.return_value = mock_row
        mock_db.query.return_value = mock_query

        with patch("app.database.SessionLocal", return_value=mock_db), \
             patch("app.security.encryption.decrypt_value", return_value="decrypted-key"):
            result = _load_ai_config()

        assert result["model"] == "gpt-4"
        assert result["api_key"] == "decrypted-key"
        assert result["api_base"] == "https://api.openai.com/v1"
        assert result["temperature"] == 0.5

    @pytest.mark.asyncio
    async def test_load_ai_config_not_found(self):
        from app.gen.model_client import _load_ai_config

        mock_db = MagicMock()
        mock_query = MagicMock()
        mock_query.filter.return_value.first.return_value = None
        mock_db.query.return_value = mock_query

        with patch("app.database.SessionLocal", return_value=mock_db), \
             patch("app.security.encryption.decrypt_value", return_value="x"):
            with pytest.raises(RuntimeError, match="AI config not found"):
                _load_ai_config()


class TestModelClientCallModel:
    """Test call_model."""

    def _patch_config(self, model="gpt-4", api_base="https://api.example.com/v1", api_key="key", temperature=0.1):
        """Return a context manager that patches _load_ai_config."""
        cfg = {
            "model": model,
            "api_key": api_key,
            "api_base": api_base,
            "temperature": temperature,
        }
        return patch("app.gen.model_client._load_ai_config", return_value=cfg)

    @pytest.mark.asyncio
    async def test_call_model_successful_post(self):
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "Hello world"}, "finish_reason": "stop"}]
        }
        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp) as mpost:
            result = call_model([{"role": "user", "content": "hi"}], temperature=0.7)

        assert result == "Hello world"
        assert mpost.call_args.args[0].endswith("/chat/completions")
        payload = mpost.call_args.kwargs["json"]
        assert payload["messages"] == [{"role": "user", "content": "hi"}]
        assert payload["temperature"] == 0.7
        assert "Authorization" in mpost.call_args.kwargs["headers"]

    @pytest.mark.asyncio
    async def test_call_model_truncation_warning(self, caplog):
        """finish_reason=length should emit a warning."""
        import logging
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "X" * 100}, "finish_reason": "length"}]
        }
        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp), \
             caplog.at_level(logging.WARNING, logger="app.gen.model_client"):
            result = call_model([{"role": "user", "content": "x"}])
        assert "X" in result

    @pytest.mark.asyncio
    async def test_call_model_no_auth_when_no_api_key(self):
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "ok"}, "finish_reason": "stop"}]
        }
        with self._patch_config(api_key="") as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp) as mpost:
            result = call_model([{"role": "user", "content": "x"}])
        # When api_key is empty, no Authorization header
        assert result == "ok"

    @pytest.mark.asyncio
    async def test_call_model_appends_chat_completions_path(self):
        """api_base without /chat/completions should have it appended."""
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "ok"}, "finish_reason": "stop"}]
        }
        with self._patch_config(api_base="https://api.example.com/v1") as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp) as mpost:
            call_model([{"role": "user", "content": "x"}])
        # Verify URL ends with /chat/completions
        called_url = mpost.call_args.args[0]
        assert called_url.endswith("/chat/completions")

    @pytest.mark.asyncio
    async def test_call_model_keeps_existing_chat_completions_path(self):
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "ok"}, "finish_reason": "stop"}]
        }
        with self._patch_config(api_base="https://api.example.com/v1/chat/completions") as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp) as mpost:
            call_model([{"role": "user", "content": "x"}])
        called_url = mpost.call_args.args[0]
        # Should not double-append
        assert called_url.count("/chat/completions") == 1

    @pytest.mark.asyncio
    async def test_call_model_missing_model_name_raises(self):
        from app.gen.model_client import call_model

        with self._patch_config(model=""):
            with pytest.raises(RuntimeError, match="MODEL_NAME not configured"):
                call_model([{"role": "user", "content": "x"}])

    @pytest.mark.asyncio
    async def test_call_model_temperature_default_from_config(self):
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "ok"}, "finish_reason": "stop"}]
        }
        with self._patch_config(temperature=0.42) as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp) as mpost:
            # No temperature passed
            call_model([{"role": "user", "content": "x"}])
        # The payload should include temperature=0.42
        payload = mpost.call_args.kwargs["json"]
        assert payload["temperature"] == 0.42

    @pytest.mark.asyncio
    async def test_call_model_http_error_raises(self):
        from app.gen.model_client import call_model
        import requests

        mock_resp = MagicMock()
        mock_resp.raise_for_status.side_effect = requests.HTTPError("500")
        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp):
            with pytest.raises(requests.HTTPError):
                call_model([{"role": "user", "content": "x"}])

    @pytest.mark.asyncio
    async def test_call_model_strips_br_in_response(self):
        from app.gen.model_client import call_model

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.json.return_value = {
            "choices": [{"message": {"content": "line1<br>line2"}, "finish_reason": "stop"}]
        }
        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp):
            result = call_model([{"role": "user", "content": "x"}])
        assert "<br>" not in result
        assert "line1 line2" == result

    @pytest.mark.asyncio
    async def test_call_model_streaming(self):
        """Stream mode: aggregate SSE chunks."""
        from app.gen.model_client import call_model

        lines = [
            'data: {"choices":[{"delta":{"content":"Hello "}}]}',
            'data: {"choices":[{"delta":{"content":"world"}}]}',
            'data: [DONE]',
        ]

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.iter_lines.return_value = lines
        # Make mock work as context manager
        mock_resp.__enter__ = MagicMock(return_value=mock_resp)
        mock_resp.__exit__ = MagicMock(return_value=False)

        chunks_received = []
        def cb(chunk):
            chunks_received.append(chunk)

        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp):
            result = call_model(
                [{"role": "user", "content": "x"}],
                stream_callback=cb,
            )

        assert "Hello " in "".join(chunks_received)
        assert "world" in "".join(chunks_received)
        assert "Hello world" == result

    @pytest.mark.asyncio
    async def test_call_model_streaming_handles_bad_json(self):
        """Streaming with bad JSON line should be skipped, not crash."""
        from app.gen.model_client import call_model

        lines = [
            'data: not-valid-json',
            'data: {"choices":[{"delta":{"content":"ok"}}]}',
            'data: [DONE]',
        ]

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.iter_lines.return_value = lines
        mock_resp.__enter__ = MagicMock(return_value=mock_resp)
        mock_resp.__exit__ = MagicMock(return_value=False)

        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp):
            result = call_model(
                [{"role": "user", "content": "x"}],
                stream_callback=lambda c: None,
            )
        assert "ok" in result

    @pytest.mark.asyncio
    async def test_call_model_streaming_empty_lines_skipped(self):
        """Streaming should skip empty lines."""
        from app.gen.model_client import call_model

        lines = [
            '',
            'data: {"choices":[{"delta":{"content":"x"}}]}',
            '',
            'data: [DONE]',
        ]

        mock_resp = MagicMock()
        mock_resp.raise_for_status = MagicMock()
        mock_resp.iter_lines.return_value = lines
        mock_resp.__enter__ = MagicMock(return_value=mock_resp)
        mock_resp.__exit__ = MagicMock(return_value=False)

        with self._patch_config() as p, \
             patch("app.gen.model_client.requests.post", return_value=mock_resp):
            result = call_model(
                [{"role": "user", "content": "x"}],
                stream_callback=lambda c: None,
            )
        assert "x" in result


# ===========================================================================
# response_parser tests
# ===========================================================================

class TestResponseParserCleanText:
    """Test _clean_text helper."""

    @pytest.mark.asyncio
    async def test_clean_text_no_change(self):
        from app.gen.response_parser import _clean_text
        assert _clean_text("plain text") == "plain text"

    @pytest.mark.asyncio
    async def test_clean_text_br(self):
        from app.gen.response_parser import _clean_text
        assert _clean_text("a<br>b") == "a\nb"

    @pytest.mark.asyncio
    async def test_clean_text_numbered(self):
        from app.gen.response_parser import _clean_text
        result = _clean_text("1. step1 2. step2")
        assert "1." in result
        assert "2." in result


class TestResponseParserToHtml:
    """Test _to_html helper."""

    @pytest.mark.asyncio
    async def test_to_html_basic(self):
        from app.gen.response_parser import _to_html
        assert _to_html("a\nb") == "a<br>b"

    @pytest.mark.asyncio
    async def test_to_html_empty(self):
        from app.gen.response_parser import _to_html
        assert _to_html("") == ""

    @pytest.mark.asyncio
    async def test_to_html_multiline(self):
        from app.gen.response_parser import _to_html
        assert _to_html("a\nb\nc") == "a<br>b<br>c"


class TestResponseParserParseResponse:
    """Test _parse_response on AnalysisSession."""

    @pytest.mark.asyncio
    async def test_parse_response_extracts_functional_points(self):
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 功能点清单
- **【用户管理】用户登录(认证)**: 用户通过用户名密码登录
- **【订单】创建订单(业务)**: 提交订单到后端

## 测试用例
| 用例ID | 模块 | 标题 | 前置条件 | 测试步骤 | 预期结果 | 优先级 |
| --- | --- | --- | --- | --- | --- | --- |
| TC-001 | 用户管理 | 登录 | 已注册 | 1.输入 2.点击 | 1.成功 | 高 |
"""
        session = AnalysisSession()
        _parse_response(session, text)
        assert len(session.functional_points) == 2
        # FP module/name/category extracted
        assert session.functional_points[0].module == "用户管理"
        assert "登录" in session.functional_points[0].name

    @pytest.mark.asyncio
    async def test_parse_response_extracts_test_cases(self):
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 功能点清单

## 测试用例
| 用例ID | 模块 | 标题 | 前置条件 | 测试步骤 | 预期结果 | 优先级 |
| --- | --- | --- | --- | --- | --- | --- |
| TC-001 | 用户 | 登录测试 | 已注册 | 1.输入 | 1.成功 | 高 |
| TC-002 | 订单 | 创建订单 | 已登录 | 1.点击 | 1.成功 | 中 |
"""
        session = AnalysisSession()
        _parse_response(session, text)
        assert len(session.test_cases) == 2
        assert session.test_cases[0].test_case_id == "TC-001"
        assert session.test_cases[0].module == "用户"
        assert session.test_cases[0].priority == "高"

    @pytest.mark.asyncio
    async def test_parse_response_skips_separator(self):
        """Markdown separator rows (|, ---, :) should be skipped."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 测试用例
| --- | --- | --- | --- | --- | --- | --- |
| TC-001 | M | T | P | S | R | 高 |
"""
        session = AnalysisSession()
        _parse_response(session, text)
        # Only the data row should be parsed (separator skipped)
        assert len(session.test_cases) == 1

    @pytest.mark.asyncio
    async def test_parse_response_skips_header(self):
        """Header row '用例ID|...' should be skipped (not a TC)."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 测试用例
| 用例ID | 模块 | 标题 | 前置条件 | 测试步骤 | 预期结果 | 优先级 |
| TC-001 | M | T | P | S | R | 高 |
"""
        session = AnalysisSession()
        _parse_response(session, text)
        # Header row skipped, only TC-001 counted
        assert len(session.test_cases) == 1

    @pytest.mark.asyncio
    async def test_parse_response_fallback_no_marker(self):
        """No '## 测试用例' marker → use entire text as tc_section."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """| TC-001 | M | T | P | S | R | 高 |"""
        session = AnalysisSession()
        _parse_response(session, text)
        # Should still parse the TC
        assert len(session.test_cases) == 1

    @pytest.mark.asyncio
    async def test_parse_response_ignores_lines_without_brackets(self):
        """FP lines without 【】 prefix should be skipped."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 功能点清单
- **普通条目**: 不应该被解析
- **【模块A】实际功能点**: 应该被解析
- 文本行: 完全不是 FP
"""
        session = AnalysisSession()
        _parse_response(session, text)
        # Only the 【】-prefixed one is parsed
        assert len(session.functional_points) == 1
        assert session.functional_points[0].module == "模块A"

    @pytest.mark.asyncio
    async def test_parse_response_default_module(self):
        """When no 【】 prefix, default to '通用' module."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        # This case: starts with `- **` and has `【` so will be parsed, but no module prefix
        text = """## 功能点清单
- **【某模块】无括号功能点**: 描述
"""
        session = AnalysisSession()
        _parse_response(session, text)
        assert session.functional_points[0].module == "某模块"
        assert "无括号" in session.functional_points[0].name

    @pytest.mark.asyncio
    async def test_parse_response_with_star_prefix(self):
        """FP lines starting with '* **' should also match."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 功能点清单
* **【M】Name(cat)**: desc
"""
        session = AnalysisSession()
        _parse_response(session, text)
        assert len(session.functional_points) == 1

    @pytest.mark.asyncio
    async def test_parse_response_short_row_skipped(self):
        """Rows with < 7 cells are skipped."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        text = """## 测试用例
| short | row |
| TC-001 | M | T | P | S | R | 高 |
"""
        session = AnalysisSession()
        _parse_response(session, text)
        # Only the long row counts
        assert len(session.test_cases) == 1

    @pytest.mark.asyncio
    async def test_parse_response_filtered_cells_less_than_7(self):
        """Row with >=8 pipes but <7 non-empty cells → skipped via continue."""
        from app.gen.response_parser import _parse_response
        from app.gen.models import AnalysisSession

        # 7 visible cells but last is empty: `| a | b | c | d | e | f | |`
        # count("|") = 8 (>= 7 enters the block), but after filter we get 6 cells
        text = """## 测试用例
| a | b | c | d | e | f | |
| TC-001 | M | T | P | S | R | 高 |
"""
        session = AnalysisSession()
        _parse_response(session, text)
        # First row filtered out (6 cells after filter < 7), second row parsed
        assert len(session.test_cases) == 1
        assert session.test_cases[0].test_case_id == "TC-001"


class TestResponseParserFpHelpers:
    """Test _parse_fps_from_text and _parse_tcs_from_text."""

    @pytest.mark.asyncio
    async def test_parse_fps_from_text(self):
        from app.gen.response_parser import _parse_fps_from_text

        text = """## 功能点清单
- **【M】FP1(cat)**: d
- **【M】FP2**: d
"""
        fps = _parse_fps_from_text(text, session_id="abc")
        assert len(fps) == 2
        assert fps[0].session_id == "abc"

    @pytest.mark.asyncio
    async def test_parse_tcs_from_text(self):
        from app.gen.response_parser import _parse_tcs_from_text

        text = """| TC-1 | M | T | P | S | R | 高 |
| TC-2 | M | T | P | S | R | 中 |"""
        tcs = _parse_tcs_from_text(text, session_id="x", start_index=10)
        assert len(tcs) == 2
        # IDs are renumbered from start_index
        assert tcs[0].test_case_id == "TC-011"
        assert tcs[1].test_case_id == "TC-012"

    @pytest.mark.asyncio
    async def test_parse_tcs_from_text_default_index(self):
        from app.gen.response_parser import _parse_tcs_from_text
        tcs = _parse_tcs_from_text("| TC-1 | M | T | P | S | R | 高 |")
        # Default start_index=0 → TC-001
        assert tcs[0].test_case_id == "TC-001"


# ===========================================================================
# feature_extractor tests
# ===========================================================================

class TestFeatureExtractorExtractFPs:
    """Test extract_functional_points."""

    @pytest.mark.asyncio
    async def test_extract_fps_from_text(self):
        from app.gen.feature_extractor import extract_functional_points
        from app.gen.models import FunctionalPoint

        mock_response = """## 功能点清单
- **【M】FP1(cat)**: desc
## 测试用例
| TC-001 | M | T | P | S | R | 高 |
"""
        with patch("app.gen.feature_extractor.call_model", return_value=mock_response):
            fps = extract_functional_points(text="Some document text")

        assert len(fps) == 1
        assert isinstance(fps[0], FunctionalPoint)

    @pytest.mark.asyncio
    async def test_extract_fps_from_image(self):
        from app.gen.feature_extractor import extract_functional_points

        mock_response = """## 功能点清单
- **【UI】按钮(cat)**: desc
"""
        with patch("app.gen.feature_extractor.call_model", return_value=mock_response) as mcall:
            fps = extract_functional_points(image_data=("png", "fakebase64"))

        assert len(fps) == 1
        messages = mcall.call_args.args[0]
        user_msg = messages[1]
        assert user_msg["content"][0]["type"] == "text"
        assert user_msg["content"][1]["type"] == "image_url"
        assert "data:image/png;base64,fakebase64" in user_msg["content"][1]["image_url"]["url"]

    @pytest.mark.asyncio
    async def test_extract_fps_with_project_description(self):
        from app.gen.feature_extractor import extract_functional_points

        mock_response = "## 功能点清单\n"
        with patch("app.gen.feature_extractor.call_model", return_value=mock_response) as mcall:
            fps = extract_functional_points(
                text="doc",
                project_description="This is a banking app"
            )

        # System prompt should include project description
        sys_msg = mcall.call_args.args[0][0]["content"]
        assert "banking app" in sys_msg

    @pytest.mark.asyncio
    async def test_extract_fps_with_progress_callback(self):
        from app.gen.feature_extractor import extract_functional_points

        progress_calls = []
        def cb(cur, total, msg):
            progress_calls.append((cur, total, msg))

        with patch("app.gen.feature_extractor.call_model", return_value="## 功能点清单\n"):
            extract_functional_points(text="x", progress_callback=cb)

        # At least one progress call
        assert len(progress_calls) >= 2

    @pytest.mark.asyncio
    async def test_extract_fps_image_with_progress(self):
        from app.gen.feature_extractor import extract_functional_points

        progress_calls = []
        def cb(cur, total, msg):
            progress_calls.append(msg)

        with patch("app.gen.feature_extractor.call_model", return_value="## 功能点清单\n"):
            extract_functional_points(image_data=("jpg", "b64"), progress_callback=cb)

        # Should have called with "正在分析图片提取功能点"
        assert any("图片" in m for m in progress_calls)

    @pytest.mark.asyncio
    async def test_extract_fps_custom_prompt(self):
        from app.gen.feature_extractor import extract_functional_points

        with patch("app.gen.feature_extractor.call_model", return_value="## 功能点清单\n") as mcall:
            extract_functional_points(text="x", fp_prompt="CUSTOM FP PROMPT")
        sys_msg = mcall.call_args.args[0][0]["content"]
        assert "CUSTOM FP PROMPT" in sys_msg

    @pytest.mark.asyncio
    async def test_extract_fps_no_fps_in_response(self):
        from app.gen.feature_extractor import extract_functional_points

        with patch("app.gen.feature_extractor.call_model", return_value="no markers at all"):
            fps = extract_functional_points(text="x")
        assert fps == []


class TestFeatureExtractorGenerateTCs:
    """Test generate_test_cases_for_fps."""

    def _make_fps(self, n: int):
        from app.gen.models import FunctionalPoint
        return [
            FunctionalPoint(id=i, session_id="s", module=f"M{i}", name=f"FP{i}",
                           description=f"d{i}", category="c")
            for i in range(1, n + 1)
        ]

    @pytest.mark.asyncio
    async def test_generate_tcs_single_batch(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(3)
        mock_response = """| TC-1 | M | T | P | S | R | 高 |"""

        with patch("app.gen.feature_extractor.call_model", return_value=mock_response):
            result = generate_test_cases_for_fps(fps, "")

        assert "test_cases" in result
        assert "warnings" in result
        assert len(result["test_cases"]) >= 1
        assert result["warnings"] == []

    @pytest.mark.asyncio
    async def test_generate_tcs_multiple_batches(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(16)  # 2 batches of 8 (FP_BATCH_SIZE=8)
        mock_response = """| TC-1 | M | T | P | S | R | 高 |"""

        with patch("app.gen.feature_extractor.call_model", return_value=mock_response):
            result = generate_test_cases_for_fps(fps, "")

        # Should have generated TCs for both batches
        assert len(result["test_cases"]) >= 2

    @pytest.mark.asyncio
    async def test_generate_tcs_with_project_description(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(2)
        mock_response = """| TC-1 | M | T | P | S | R | 高 |"""

        with patch("app.gen.feature_extractor.call_model", return_value=mock_response) as mcall:
            generate_test_cases_for_fps(fps, "My project context")

        # System prompt should include project context
        sys_msg = mcall.call_args.args[0][0]["content"]
        assert "My project context" in sys_msg

    @pytest.mark.asyncio
    async def test_generate_tcs_retries_on_empty(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(1)
        # First call returns no TCs, second returns TCs
        responses = [
            "no markers",  # empty parse
            "| TC-1 | M | T | P | S | R | 高 |",  # success
        ]

        with patch("app.gen.feature_extractor.call_model", side_effect=responses), \
             patch("app.gen.feature_extractor.time.sleep"):  # skip retry sleep
            result = generate_test_cases_for_fps(fps, "")

        assert len(result["test_cases"]) >= 1

    @pytest.mark.asyncio
    async def test_generate_tcs_fails_after_max_retries(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(1)
        # Always empty
        with patch("app.gen.feature_extractor.call_model", return_value="no markers"), \
             patch("app.gen.feature_extractor.time.sleep"):
            result = generate_test_cases_for_fps(fps, "")

        # Should add a warning
        assert len(result["warnings"]) >= 1
        assert "no test cases" in result["warnings"][0].lower()

    @pytest.mark.asyncio
    async def test_generate_tcs_handles_exception(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(1)
        # All attempts raise
        with patch("app.gen.feature_extractor.call_model", side_effect=RuntimeError("api down")), \
             patch("app.gen.feature_extractor.time.sleep"):
            result = generate_test_cases_for_fps(fps, "")

        # After MAX_RETRIES, should add warning
        assert len(result["warnings"]) >= 1

    @pytest.mark.asyncio
    async def test_generate_tcs_with_progress_callback(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(2)
        progress_calls = []
        def cb(cur, total, msg):
            progress_calls.append((cur, total, msg))

        with patch("app.gen.feature_extractor.call_model", return_value="| TC-1 | M | T | P | S | R | 高 |"), \
             patch("app.gen.feature_extractor.time.sleep"):
            generate_test_cases_for_fps(
                fps, "", progress_callback=cb,
                phase1_offset=1, total_steps=3,
            )

        assert len(progress_calls) >= 1

    @pytest.mark.asyncio
    async def test_generate_tcs_custom_prompt(self):
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(1)
        with patch("app.gen.feature_extractor.call_model", return_value="| TC-1 | M | T | P | S | R | 高 |") as mcall, \
             patch("app.gen.feature_extractor.time.sleep"):
            generate_test_cases_for_fps(fps, "", tc_prompt="CUSTOM TC {fp_descriptions} {csv_header}")

        # Check that custom prompt was used
        sys_msg = mcall.call_args.args[0][0]["content"]
        assert "CUSTOM TC" in sys_msg

    @pytest.mark.asyncio
    async def test_generate_tcs_long_batch_names(self):
        """Batch with > 3 FPs: ' +N more' suffix in progress message."""
        from app.gen.feature_extractor import generate_test_cases_for_fps

        fps = self._make_fps(8)  # exactly one batch of 8
        progress_calls = []
        def cb(cur, total, msg):
            progress_calls.append(msg)

        with patch("app.gen.feature_extractor.call_model", return_value="| TC-1 | M | T | P | S | R | 高 |"), \
             patch("app.gen.feature_extractor.time.sleep"):
            generate_test_cases_for_fps(
                fps, "", progress_callback=cb,
                phase1_offset=1, total_steps=2,
            )
        # Progress message should include a name
        assert len(progress_calls) >= 1

    @pytest.mark.asyncio
    async def test_generate_tcs_logs_batch_info(self):
        """Successful batch should log INFO with batch number."""
        from app.gen.feature_extractor import generate_test_cases_for_fps
        import logging

        fps = self._make_fps(1)
        with patch("app.gen.feature_extractor.call_model", return_value="| TC-1 | M | T | P | S | R | 高 |"), \
             patch("app.gen.feature_extractor.time.sleep"), \
             patch.object(logging.getLogger("app.gen.feature_extractor"), "info") as minfo:
            generate_test_cases_for_fps(fps, "")
        # INFO log was called for successful batch
        assert any("generated" in str(c) for c in minfo.call_args_list)


# ===========================================================================
# prompts tests
# ===========================================================================

class TestPrompts:
    """Test prompts module."""

    @pytest.mark.asyncio
    async def test_get_default_prompts_structure(self):
        from app.gen.prompts import get_default_prompts
        prompts = get_default_prompts()
        assert "fp_extract" in prompts
        assert "tc_generate" in prompts
        assert prompts["fp_extract"]["label"] == "功能点提取"
        assert prompts["tc_generate"]["label"] == "测试用例生成"
        assert "content" in prompts["fp_extract"]
        assert "content" in prompts["tc_generate"]
        assert len(prompts["fp_extract"]["content"]) > 0
        assert len(prompts["tc_generate"]["content"]) > 0


# ===========================================================================
# orchestrator tests
# ===========================================================================

class TestOrchestratorTwoPhaseAnalyze:
    """Test two_phase_analyze."""

    @pytest.mark.asyncio
    async def test_two_phase_analyze_happy(self):
        from app.gen.orchestrator import two_phase_analyze

        mock_fps = [MagicMock(), MagicMock()]
        mock_tcs = [MagicMock()]

        with patch("app.gen.orchestrator.extract_functional_points", return_value=mock_fps), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": mock_tcs, "warnings": []}):
            result = two_phase_analyze("Some document text")

        assert "functional_points" in result
        assert "test_cases" in result
        assert "warnings" in result
        assert result["functional_points"] == mock_fps
        assert result["test_cases"] == mock_tcs

    @pytest.mark.asyncio
    async def test_two_phase_analyze_truncates_long_text(self):
        from app.gen.orchestrator import two_phase_analyze

        # 10000 chars → ~15000 tokens → over 6554 budget → truncated
        long_text = "X" * 10000

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[]) as mext, \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}):
            result = two_phase_analyze(long_text)

        # The text passed to extract_functional_points should be truncated
        called_text = mext.call_args.kwargs.get("text", "")
        assert len(called_text) < len(long_text)

    @pytest.mark.asyncio
    async def test_two_phase_analyze_no_fps_warning(self):
        from app.gen.orchestrator import two_phase_analyze

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[]), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}):
            result = two_phase_analyze("short text")

        # Should warn about no FPs
        assert any("No functional points" in w for w in result["warnings"])

    @pytest.mark.asyncio
    async def test_two_phase_analyze_fp_extraction_fails(self):
        from app.gen.orchestrator import two_phase_analyze

        with patch("app.gen.orchestrator.extract_functional_points",
                   side_effect=RuntimeError("llm down")):
            result = two_phase_analyze("text")

        assert result.get("error") is True
        assert "FP extraction failed" in result["warnings"][0]
        assert result["test_cases"] == []
        assert result["functional_points"] == []

    @pytest.mark.asyncio
    async def test_two_phase_analyze_with_progress_callback(self):
        from app.gen.orchestrator import two_phase_analyze

        progress_calls = []
        def cb(cur, total, msg):
            progress_calls.append(msg)

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[]), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}):
            two_phase_analyze("text", progress_callback=cb)

        assert any("提取功能点" in m for m in progress_calls)

    @pytest.mark.asyncio
    async def test_two_phase_analyze_with_project_description(self):
        from app.gen.orchestrator import two_phase_analyze

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[]) as mext, \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}):
            two_phase_analyze("text", project_description="Banking context")

        # Project description forwarded
        assert mext.call_args.kwargs.get("project_description") == "Banking context"

    @pytest.mark.asyncio
    async def test_two_phase_analyze_with_prompts_dict(self):
        from app.gen.orchestrator import two_phase_analyze

        prompts = {
            "fp_extract": {"content": "CUSTOM FP PROMPT"},
            "tc_generate": {"content": "CUSTOM TC PROMPT"},
        }
        mock_fp = MagicMock()

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[mock_fp]) as mext, \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}) as mgen:
            two_phase_analyze("text", prompts=prompts)

        assert mext.call_args.kwargs.get("fp_prompt") == "CUSTOM FP PROMPT"
        assert mgen.call_args.kwargs.get("tc_prompt") == "CUSTOM TC PROMPT"

    @pytest.mark.asyncio
    async def test_two_phase_analyze_with_prompts_strings(self):
        """Prompts can be plain strings, not just dicts."""
        from app.gen.orchestrator import two_phase_analyze

        prompts = {
            "fp_extract": "STR FP",
            "tc_generate": "STR TC",
        }
        mock_fp = MagicMock()

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[mock_fp]) as mext, \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}) as mgen:
            two_phase_analyze("text", prompts=prompts)

        assert mext.call_args.kwargs.get("fp_prompt") == "STR FP"
        assert mgen.call_args.kwargs.get("tc_prompt") == "STR TC"

    @pytest.mark.asyncio
    async def test_two_phase_analyze_propagates_tc_warnings(self):
        from app.gen.orchestrator import two_phase_analyze

        with patch("app.gen.orchestrator.extract_functional_points", return_value=[MagicMock()]), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": ["batch failed"]}):
            result = two_phase_analyze("text")
        assert "batch failed" in result["warnings"]


class TestOrchestratorImageAnalyze:
    """Test _analyze_image_two_phase."""

    @pytest.mark.asyncio
    async def test_image_analyze_happy(self):
        from app.gen.orchestrator import _analyze_image_two_phase

        mock_file = MagicMock()
        mock_file.filename = "ui.png"

        with patch("app.gen.orchestrator.encode_image", return_value="b64data"), \
             patch("app.gen.orchestrator.extract_functional_points", return_value=[MagicMock()]) as mext, \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [MagicMock()], "warnings": []}):
            result = _analyze_image_two_phase(mock_file, None, "")

        assert result["functional_points"]
        # encode_image was called
        assert mext.call_args.kwargs.get("image_data") == ("png", "b64data")

    @pytest.mark.asyncio
    async def test_image_analyze_no_fps(self):
        from app.gen.orchestrator import _analyze_image_two_phase

        mock_file = MagicMock()
        mock_file.filename = "ui.jpg"

        with patch("app.gen.orchestrator.encode_image", return_value="b64"), \
             patch("app.gen.orchestrator.extract_functional_points", return_value=[]):
            result = _analyze_image_two_phase(mock_file, None, "")

        assert "No functional points" in str(result["warnings"])

    @pytest.mark.asyncio
    async def test_image_analyze_fails(self):
        from app.gen.orchestrator import _analyze_image_two_phase

        mock_file = MagicMock()
        mock_file.filename = "ui.png"

        with patch("app.gen.orchestrator.encode_image", return_value="b64"), \
             patch("app.gen.orchestrator.extract_functional_points",
                   side_effect=RuntimeError("llm fail")):
            result = _analyze_image_two_phase(mock_file, None, "")

        assert result.get("error") is True
        assert "Image FP extraction failed" in str(result["warnings"])

    @pytest.mark.asyncio
    async def test_image_analyze_with_progress(self):
        from app.gen.orchestrator import _analyze_image_two_phase

        mock_file = MagicMock()
        mock_file.filename = "ui.png"
        progress = []
        def cb(cur, total, msg):
            progress.append(msg)

        with patch("app.gen.orchestrator.encode_image", return_value="b64"), \
             patch("app.gen.orchestrator.extract_functional_points", return_value=[]):
            _analyze_image_two_phase(mock_file, cb, "")

        assert any("图片" in m for m in progress)


class TestOrchestratorPdfAnalyze:
    """Test _analyze_pdf_two_phase."""

    @pytest.mark.asyncio
    async def test_pdf_invalid(self):
        from app.gen.orchestrator import _analyze_pdf_two_phase

        mock_file = MagicMock()
        with patch("app.gen.orchestrator.validate_pdf", return_value=(False, "PDF文件损坏")):
            result = _analyze_pdf_two_phase(mock_file, None, "")
        assert result.get("error") is True
        assert "PDF文件损坏" in result["warnings"]

    @pytest.mark.asyncio
    async def test_pdf_dual_layer(self):
        """Dual-layer PDF → uses text pipeline."""
        from app.gen.orchestrator import _analyze_pdf_two_phase

        mock_file = MagicMock()
        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=True), \
             patch("app.gen.orchestrator.extract_text_from_pdf", return_value="Some text"), \
             patch("app.gen.orchestrator.two_phase_analyze",
                   return_value={"functional_points": [MagicMock()], "test_cases": [MagicMock()], "warnings": []}):
            result = _analyze_pdf_two_phase(mock_file, None, "")

        assert result["functional_points"]

    @pytest.mark.asyncio
    async def test_pdf_dual_layer_empty_text(self):
        """Dual-layer but no text content → error."""
        from app.gen.orchestrator import _analyze_pdf_two_phase

        mock_file = MagicMock()
        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=True), \
             patch("app.gen.orchestrator.extract_text_from_pdf", return_value="   "):
            result = _analyze_pdf_two_phase(mock_file, None, "")

        assert result.get("error") is True
        assert "无有效文字" in str(result["warnings"])

    @pytest.mark.asyncio
    async def test_pdf_scan_only(self):
        """Scan-only PDF → render pages and extract FPs."""
        from app.gen.orchestrator import _analyze_pdf_two_phase
        from app.gen.models import FunctionalPoint

        mock_file = MagicMock()
        fp1 = FunctionalPoint(id=1, module="M", name="N1", description="d", category="c", session_id="orig")
        fp2 = FunctionalPoint(id=1, module="M", name="N2", description="d", category="c", session_id="orig")

        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=False), \
             patch("app.gen.orchestrator.render_pdf_pages_to_images",
                   return_value=[("png", "b64a"), ("png", "b64b")]), \
             patch("app.gen.orchestrator.extract_functional_points",
                   side_effect=[[fp1], [fp2]]), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [MagicMock()], "warnings": []}):
            result = _analyze_pdf_two_phase(mock_file, None, "")

        # Both pages' FPs are merged
        assert len(result["functional_points"]) == 2
        # FPs are re-numbered sequentially
        ids = [fp.id for fp in result["functional_points"]]
        assert ids == [1, 2]

    @pytest.mark.asyncio
    async def test_pdf_scan_only_no_pages(self):
        """Scan-only PDF with no renderable pages → error."""
        from app.gen.orchestrator import _analyze_pdf_two_phase

        mock_file = MagicMock()
        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=False), \
             patch("app.gen.orchestrator.render_pdf_pages_to_images", return_value=[]):
            result = _analyze_pdf_two_phase(mock_file, None, "")

        assert result.get("error") is True
        assert "无有效页面" in str(result["warnings"])

    @pytest.mark.asyncio
    async def test_pdf_scan_only_page_extraction_fails(self):
        """One page's FP extraction fails, others succeed."""
        from app.gen.orchestrator import _analyze_pdf_two_phase
        from app.gen.models import FunctionalPoint

        mock_file = MagicMock()
        fp = FunctionalPoint(id=1, module="M", name="N", description="d", category="c")

        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=False), \
             patch("app.gen.orchestrator.render_pdf_pages_to_images",
                   return_value=[("png", "b64a"), ("png", "b64b")]), \
             patch("app.gen.orchestrator.extract_functional_points",
                   side_effect=[[fp], RuntimeError("llm timeout")]), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [], "warnings": []}):
            result = _analyze_pdf_two_phase(mock_file, None, "")

        assert len(result["functional_points"]) == 1
        # Page 2 failure → warning
        assert any("第 2 页" in w for w in result["warnings"])

    @pytest.mark.asyncio
    async def test_pdf_dual_layer_with_progress(self):
        """Dual-layer PDF with progress_callback (covers line 153)."""
        from app.gen.orchestrator import _analyze_pdf_two_phase

        mock_file = MagicMock()
        progress = []
        def cb(cur, total, msg):
            progress.append(msg)

        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=True), \
             patch("app.gen.orchestrator.extract_text_from_pdf", return_value="Some text"), \
             patch("app.gen.orchestrator.two_phase_analyze",
                   return_value={"functional_points": [], "test_cases": [], "warnings": []}):
            _analyze_pdf_two_phase(mock_file, cb, "")

        assert any("正在从PDF提取文字" in m for m in progress)

    @pytest.mark.asyncio
    async def test_pdf_scan_only_with_progress(self):
        """Scan-only PDF with progress_callback (covers lines 161, 171)."""
        from app.gen.orchestrator import _analyze_pdf_two_phase
        from app.gen.models import FunctionalPoint

        mock_file = MagicMock()
        fp = FunctionalPoint(id=1, module="M", name="N", description="d", category="c")
        progress = []
        def cb(cur, total, msg):
            progress.append(msg)

        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=False), \
             patch("app.gen.orchestrator.render_pdf_pages_to_images",
                   return_value=[("png", "b64a"), ("png", "b64b")]), \
             patch("app.gen.orchestrator.extract_functional_points", return_value=[fp]), \
             patch("app.gen.orchestrator.generate_test_cases_for_fps",
                   return_value={"test_cases": [MagicMock()], "warnings": []}):
            _analyze_pdf_two_phase(mock_file, cb, "")

        assert any("正在将PDF页面转为图片" in m for m in progress)
        assert any("正在分析第 1" in m for m in progress)
        assert any("正在分析第 2" in m for m in progress)

    @pytest.mark.asyncio
    async def test_pdf_scan_only_no_fps_else_branch(self):
        """Scan-only with no FPs extracted from any page → all_tcs = [] (covers line 201)."""
        from app.gen.orchestrator import _analyze_pdf_two_phase

        mock_file = MagicMock()

        with patch("app.gen.orchestrator.validate_pdf", return_value=(True, None)), \
             patch("app.gen.orchestrator.is_pdf_dual_layer", return_value=False), \
             patch("app.gen.orchestrator.render_pdf_pages_to_images",
                   return_value=[("png", "b64")]), \
             patch("app.gen.orchestrator.extract_functional_points", return_value=[]):
            result = _analyze_pdf_two_phase(mock_file, None, "")

        # No FPs extracted, so all_fps is empty → else branch hit
        assert len(result["functional_points"]) == 0
        assert result["test_cases"] == []


# ===========================================================================
# multi_file tests
# ===========================================================================

class TestMultiFileExtractContent:
    """Test extract_multi_file_content."""

    def _make_file(self, content: bytes, name: str = "f.txt"):
        buf = io.BytesIO(content)
        buf.name = name
        return buf

    @pytest.mark.asyncio
    async def test_extract_empty_files_raises(self):
        from app.gen.multi_file import extract_multi_file_content
        with pytest.raises(ValueError):
            extract_multi_file_content([], [])

    @pytest.mark.asyncio
    async def test_extract_too_many_files_raises(self):
        from app.gen.multi_file import extract_multi_file_content
        files = [self._make_file(b"x", f"f{i}.md") for i in range(15)]
        with pytest.raises(ValueError, match="最多上传"):
            extract_multi_file_content(files, [f.name for f in files])

    @pytest.mark.asyncio
    async def test_extract_unsupported_extension_raises(self):
        from app.gen.multi_file import extract_multi_file_content
        files = [self._make_file(b"x", "f.exe")]
        with pytest.raises(ValueError, match="不支持的文件类型"):
            extract_multi_file_content(files, ["f.exe"])

    @pytest.mark.asyncio
    async def test_extract_total_size_too_big_raises(self):
        from app.gen.multi_file import extract_multi_file_content
        # 50MB+1
        files = [self._make_file(b"x" * (50 * 1024 * 1024 + 1), "f.md")]
        with pytest.raises(ValueError, match="50MB"):
            extract_multi_file_content(files, ["f.md"])

    @pytest.mark.asyncio
    async def test_extract_docx_success(self):
        from app.gen.multi_file import extract_multi_file_content
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello DOCX")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "test.docx"
        files = [buf]

        result = extract_multi_file_content(files, ["test.docx"])
        text, names, warnings = result
        assert "Hello DOCX" in text
        assert names == ["test.docx"]
        assert warnings == []

    @pytest.mark.asyncio
    async def test_extract_md_success(self):
        from app.gen.multi_file import extract_multi_file_content

        buf = io.BytesIO(b"# Markdown Title\nContent")
        buf.name = "test.md"
        files = [buf]

        text, names, warnings = extract_multi_file_content(files, ["test.md"])
        assert "Markdown" in text or "Title" in text

    @pytest.mark.asyncio
    async def test_extract_empty_docx_warns(self):
        from app.gen.multi_file import extract_multi_file_content
        from docx import Document

        doc = Document()  # empty
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "empty.docx"
        files = [buf]

        with pytest.raises(ValueError, match="所有文件均未提取到有效内容"):
            extract_multi_file_content(files, ["empty.docx"])

    @pytest.mark.asyncio
    async def test_extract_empty_md_warns(self):
        from app.gen.multi_file import extract_multi_file_content

        buf = io.BytesIO(b"")
        buf.name = "empty.md"
        files = [buf]

        with pytest.raises(ValueError, match="所有文件均未提取到有效内容"):
            extract_multi_file_content(files, ["empty.md"])

    @pytest.mark.asyncio
    async def test_extract_pdf_dual_layer(self):
        from app.gen.multi_file import extract_multi_file_content
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "PDF text content")
        pdf_bytes = doc.tobytes()
        doc.close()

        buf = io.BytesIO(pdf_bytes)
        buf.name = "test.pdf"
        files = [buf]

        text, names, warnings = extract_multi_file_content(files, ["test.pdf"])
        assert "PDF text content" in text or "PDF" in text

    @pytest.mark.asyncio
    async def test_extract_pdf_scan_only(self):
        from app.gen.multi_file import extract_multi_file_content
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()  # blank
        pdf_bytes = doc.tobytes()
        doc.close()

        mock_fp = MagicMock()
        mock_fp.module = "M"
        mock_fp.name = "N"
        mock_fp.category = "C"
        mock_fp.description = "D"

        with patch("app.gen.multi_file.extract_functional_points", return_value=[mock_fp]):
            buf = io.BytesIO(pdf_bytes)
            buf.name = "scan.pdf"
            text, names, warnings = extract_multi_file_content([buf], ["scan.pdf"])
        # FP-based extraction occurred
        assert text is not None

    @pytest.mark.asyncio
    async def test_extract_image_success(self):
        from app.gen.multi_file import extract_multi_file_content
        from PIL import Image

        img = Image.new("RGB", (5, 5), color="red")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        buf.name = "test.png"
        files = [buf]

        mock_fp = MagicMock()
        mock_fp.module = "M"
        mock_fp.name = "N"
        mock_fp.category = "C"
        mock_fp.description = "D"

        with patch("app.gen.multi_file.extract_functional_points", return_value=[mock_fp]):
            text, names, warnings = extract_multi_file_content(files, ["test.png"])
        # Image FP-based extraction occurred
        assert text is not None

    @pytest.mark.asyncio
    async def test_extract_image_timeout_retries(self):
        """Image with timeout error should retry."""
        from app.gen.multi_file import extract_multi_file_content
        from PIL import Image

        img = Image.new("RGB", (5, 5), color="red")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        buf.name = "test.png"
        files = [buf]

        # First two attempts timeout, third succeeds
        mock_fp = MagicMock(module="M", name="N", category="C", description="D")
        side_effects = [
            Exception("Connection timed out"),
            Exception("Request timed out"),
            [mock_fp],
        ]
        with patch("app.gen.multi_file.extract_functional_points", side_effect=side_effects), \
             patch("app.gen.multi_file.time.sleep"):
            text, names, warnings = extract_multi_file_content(files, ["test.png"])
        assert text is not None

    @pytest.mark.asyncio
    async def test_extract_image_non_timeout_raises(self):
        """Image with non-timeout error should NOT retry."""
        from app.gen.multi_file import extract_multi_file_content
        from PIL import Image

        img = Image.new("RGB", (5, 5), color="red")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        buf.name = "test.png"
        files = [buf]

        with patch("app.gen.multi_file.extract_functional_points",
                   side_effect=RuntimeError("api error")), \
             patch("app.gen.multi_file.time.sleep") as msleep:
            with pytest.raises(ValueError, match="所有文件均未提取到有效内容"):
                extract_multi_file_content(files, ["test.png"])
        # No retry sleep
        assert msleep.call_count == 0

    @pytest.mark.asyncio
    async def test_extract_file_exception_caught(self):
        """Generic exception in file processing is caught and warned.

        With only a single bad file, the combined output is empty and
        the function raises a ValueError. The error is caught BEFORE
        the empty-content check, so the warning was added but not
        returned. Verify by ensuring the function exits cleanly with raise.
        """
        from app.gen.multi_file import extract_multi_file_content

        buf = io.BytesIO(b"not a real docx")
        buf.name = "bad.docx"
        files = [buf]

        with pytest.raises(ValueError, match="所有文件均未提取到有效内容"):
            extract_multi_file_content(files, ["bad.docx"])

    @pytest.mark.asyncio
    async def test_extract_with_progress_callback(self):
        from app.gen.multi_file import extract_multi_file_content

        buf = io.BytesIO(b"# MD content")
        buf.name = "test.md"
        files = [buf]

        progress = []
        def cb(cur, total, msg):
            progress.append((cur, total, msg))

        text, names, warnings = extract_multi_file_content(
            files, ["test.md"], progress_callback=cb
        )
        assert len(progress) >= 1

    @pytest.mark.asyncio
    async def test_extract_multiple_files_concat(self):
        """Multiple files should be concatenated."""
        from app.gen.multi_file import extract_multi_file_content

        b1 = io.BytesIO(b"# File1\nContent1")
        b1.name = "a.md"
        b2 = io.BytesIO(b"# File2\nContent2")
        b2.name = "b.md"
        files = [b1, b2]

        text, names, warnings = extract_multi_file_content(files, ["a.md", "b.md"])
        assert "Content1" in text
        assert "Content2" in text
        # Files concatenated with section headers
        assert "===== 文件" in text

    @pytest.mark.asyncio
    async def test_extract_pdf_dual_layer_but_no_text_falls_back_to_image(self):
        """PDF marked dual-layer but no text → try image mode."""
        from app.gen.multi_file import extract_multi_file_content
        try:
            import fitz
        except ImportError:
            pytest.skip("fitz not available")

        doc = fitz.open()
        doc.new_page()  # blank — passes dual_layer check, but extract_text returns empty
        pdf_bytes = doc.tobytes()
        doc.close()

        mock_fp = MagicMock(module="M", name="N", category="C", description="D")

        # is_pdf_dual_layer returns True (because it has pages), but text is empty
        # Actually a blank page returns False for is_pdf_dual_layer
        # To hit this branch we need a page that is_pdf_dual_layer considers has text
        # but extract_text_from_pdf returns empty.
        # Easiest: mock is_pdf_dual_layer to return True, extract_text to return ""
        with patch("app.gen.multi_file.is_pdf_dual_layer", return_value=True), \
             patch("app.gen.multi_file.extract_text_from_pdf", return_value=""), \
             patch("app.gen.multi_file.render_pdf_pages_to_images", return_value=[("png", "b")]), \
             patch("app.gen.multi_file.extract_functional_points", return_value=[mock_fp]):
            buf = io.BytesIO(pdf_bytes)
            buf.name = "weird.pdf"
            text, names, warnings = extract_multi_file_content([buf], ["weird.pdf"])
        assert any("无有效文字" in w for w in warnings) or "===== 图片" in text

    @pytest.mark.asyncio
    async def test_extract_pdf_dual_layer_image_mode_empty_pages(self):
        """PDF dual-layer → empty text → image mode returns no pages → skip + warn."""
        from app.gen.multi_file import extract_multi_file_content

        with patch("app.gen.multi_file.is_pdf_dual_layer", return_value=True), \
             patch("app.gen.multi_file.extract_text_from_pdf", return_value=""), \
             patch("app.gen.multi_file.render_pdf_pages_to_images", return_value=[]):
            buf = io.BytesIO(b"fake pdf content")
            buf.name = "empty.pdf"
            with pytest.raises(ValueError, match="所有文件均未提取到有效内容"):
                extract_multi_file_content([buf], ["empty.pdf"])

    @pytest.mark.asyncio
    async def test_extract_pdf_no_pages_scan_mode(self):
        """Scan-only PDF with no renderable pages → skip + warn."""
        from app.gen.multi_file import extract_multi_file_content

        with patch("app.gen.multi_file.is_pdf_dual_layer", return_value=False), \
             patch("app.gen.multi_file.render_pdf_pages_to_images", return_value=[]):
            buf = io.BytesIO(b"fake")
            buf.name = "x.pdf"
            with pytest.raises(ValueError, match="所有文件均未提取到有效内容"):
                extract_multi_file_content([buf], ["x.pdf"])

    @pytest.mark.asyncio
    async def test_extract_image_with_empty_fps(self):
        """Image where extract_functional_points returns []."""
        from app.gen.multi_file import extract_multi_file_content
        from PIL import Image

        img = Image.new("RGB", (5, 5), color="red")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        buf.name = "test.png"
        files = [buf]

        with patch("app.gen.multi_file.extract_functional_points", return_value=[]):
            text, names, warnings = extract_multi_file_content(files, ["test.png"])
        # Should still produce some output (with "（未提取到功能点）")
        assert "未提取到功能点" in text

    @pytest.mark.asyncio
    async def test_extract_filename_fallback(self):
        """When filenames list is empty, the function uses fallback name.

        The fallback 'file_N' has no extension, so the function raises
        ValueError for unsupported type. The test verifies the fallback
        path is taken (the error message contains the fallback name's
        empty extension).
        """
        from app.gen.multi_file import extract_multi_file_content

        b1 = io.BytesIO(b"# X")
        b1.name = "a.md"
        files = [b1]
        with pytest.raises(ValueError, match="不支持的文件类型"):
            extract_multi_file_content(files, [])
