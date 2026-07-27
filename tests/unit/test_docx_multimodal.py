"""Tests for ordered DOCX parsing and multimodal content assembly."""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from app.gen.docx_parser import (
    blocks_to_plain_text,
    extract_ordered_blocks,
    extract_text,
)
from app.gen.feature_extractor import content_parts_to_openai_user_content
from app.gen.prompts import (
    FP_BATCH_SIZE,
    FP_EXTRACT_PROMPT,
    MIN_TCS_PER_ITEM,
    TC_GENERATE_PROMPT,
    TC_GENERATE_UI_PROMPT,
)


def _tiny_png_bytes() -> bytes:
    buf = BytesIO()
    Image.new("RGB", (4, 4), color=(0, 128, 255)).save(buf, format="PNG")
    return buf.getvalue()


def _docx_bytes(*, with_image: bool = True) -> bytes:
    doc = Document()
    doc.add_paragraph("前置文字AAA")
    if with_image:
        img_buf = BytesIO(_tiny_png_bytes())
        doc.add_picture(img_buf, width=Inches(0.3))
    doc.add_paragraph("后置文字BBB")
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


class TestDocxOrderedBlocks:
    def test_text_only(self):
        data = _docx_bytes(with_image=False)
        blocks, warnings = extract_ordered_blocks(BytesIO(data))
        assert warnings == []
        assert all(b["type"] == "text" for b in blocks)
        plain = extract_text(BytesIO(data))
        assert "前置文字AAA" in plain
        assert "后置文字BBB" in plain

    def test_text_image_text_order(self):
        data = _docx_bytes(with_image=True)
        blocks, warnings = extract_ordered_blocks(BytesIO(data))
        assert warnings == []
        types = [b["type"] for b in blocks]
        assert types == ["text", "image", "text"]
        assert "前置文字AAA" in blocks[0]["text"]
        assert blocks[1]["ext"] == "png"
        assert blocks[1]["b64"]
        assert "后置文字BBB" in blocks[2]["text"]

    def test_max_images_warning(self):
        data = _docx_bytes(with_image=True)
        blocks, warnings = extract_ordered_blocks(BytesIO(data), max_images=0)
        assert all(b["type"] == "text" for b in blocks)
        assert any("跳过" in w for w in warnings)

    def test_blocks_to_plain_text_placeholders(self):
        data = _docx_bytes(with_image=True)
        blocks, _ = extract_ordered_blocks(BytesIO(data))
        plain = blocks_to_plain_text(blocks)
        assert "前置文字AAA" in plain
        assert "[图片1]" in plain
        assert "后置文字BBB" in plain


class TestMultimodalAssembly:
    def test_preserves_order(self):
        parts = [
            {"type": "text", "text": "hello"},
            {"type": "image", "ext": "png", "b64": "abc123"},
            {"type": "text", "text": "world"},
        ]
        content = content_parts_to_openai_user_content(parts)
        # intro + 3 parts
        assert content[0]["type"] == "text"
        assert content[1] == {"type": "text", "text": "hello"}
        assert content[2]["type"] == "image_url"
        assert content[2]["image_url"]["url"].startswith("data:image/png;base64,abc123")
        assert content[3] == {"type": "text", "text": "world"}


class TestPromptTestItems:
    def test_fp_extract_mentions_test_items(self):
        assert "测试项" in FP_EXTRACT_PROMPT
        assert "functional_points" in FP_EXTRACT_PROMPT

    def test_tc_prompts_require_three_cases(self):
        assert "至少生成 3 条" in TC_GENERATE_PROMPT or "至少 3 条" in TC_GENERATE_PROMPT
        assert "至少生成 3 条" in TC_GENERATE_UI_PROMPT or "至少 3 条" in TC_GENERATE_UI_PROMPT
        assert MIN_TCS_PER_ITEM == 3
        assert FP_BATCH_SIZE == 3
